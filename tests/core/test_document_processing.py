import pytest
import os
from unittest.mock import patch, MagicMock, mock_open

# Import LangchainDocument for creating test data
from langchain_core.documents import Document as LangchainDocument

# Specific error for DOCX
from docx.opc.exceptions import PackageNotFoundError

# Specific error for PDF
import pdfplumber # Keep for pdfplumber.exceptions.PDFSyntaxError below if used directly
from pdfplumber.exceptions import PDFSyntaxError # Import directly

# Modules to test
from core.document_processing import (
    save_uploaded_file,
    load_document,
    chunk_documents,
    index_documents,
)

# Import config to use/mock its values, and logger for mocking
from core import config  # For PDF_STORAGE_PATH

# core.logger_config is not directly imported here, but we'll patch the logger instance
# used within document_processing.py, which is 'core.document_processing.logger'

# Paths for mocking
MOCK_OPEN_PATH = "builtins.open" # Used by save_uploaded_file and load_document (for .txt)
OS_PATH_JOIN_PATH = "core.document_processing.os.path.join" # Used by save_uploaded_file
OS_PATH_BASENAME_PATH = "core.document_processing.os.path.basename"  # Used by load_document
PDF_PLUMBER_LOADER_PATH = "core.document_processing.PDFPlumberLoader"
DOCX_MODULE_PATH = "core.document_processing.docx"  # To mock docx.Document
RECURSIVE_SPLITTER_PATH = "core.document_processing.RecursiveCharacterTextSplitter"
STREAMLIT_ERROR_PATH = "core.document_processing.st.error"
STREAMLIT_WARNING_PATH = "core.document_processing.st.warning"
DOCUMENT_PROCESSING_LOGGER_PATH = "core.document_processing.logger"


# Define the path to test data files - these files won't exist in this environment
# but tests for load_document will mock their loading.
TEST_DATA_DIR = "test_data_dummy"  # Dummy dir, actual file content will be mocked
SAMPLE_TXT_PATH = os.path.join(TEST_DATA_DIR, "sample.txt")
SAMPLE_DOCX_PATH = os.path.join(TEST_DATA_DIR, "sample.docx")
EMPTY_TXT_PATH = os.path.join(TEST_DATA_DIR, "empty.txt")
EMPTY_DOCX_PATH = os.path.join(TEST_DATA_DIR, "empty.docx")
CORRUPTED_DOCX_PATH = os.path.join(TEST_DATA_DIR, "corrupted.docx")
BAD_ENCODING_TXT_PATH = os.path.join(TEST_DATA_DIR, "bad_encoding.txt")
CORRUPTED_PDF_PATH = os.path.join(TEST_DATA_DIR, "corrupted.pdf")
UNSUPPORTED_FILE_PATH = os.path.join(TEST_DATA_DIR, "sample.png")
NON_EXISTENT_TXT_PATH = os.path.join(TEST_DATA_DIR, "non_existent.txt")


@pytest.fixture(autouse=True)
def mock_logger_fixture():
    """Automatically mock the logger for all tests in this file."""
    with patch(DOCUMENT_PROCESSING_LOGGER_PATH) as mock_log:
        yield mock_log


# --- Tests for save_uploaded_file (from previous step, confirmed good) ---


@patch(OS_PATH_JOIN_PATH, return_value="/fake/path/to/file.pdf")
@patch(MOCK_OPEN_PATH)
def test_save_uploaded_file_success(mock_open_file, mock_os_join, mock_logger_fixture):
    mock_uploaded_file = MagicMock()
    mock_uploaded_file.name = "file.pdf"
    mock_uploaded_file.getbuffer.return_value = b"file_content"

    with patch.object(config, "PDF_STORAGE_PATH", "/fake/storage"):
        result_path = save_uploaded_file(mock_uploaded_file)

    mock_os_join.assert_called_once_with("/fake/storage", "file.pdf")
    mock_open_file.assert_called_once_with("/fake/path/to/file.pdf", "wb")
    mock_open_file().write.assert_called_once_with(b"file_content")
    assert result_path == "/fake/path/to/file.pdf"
    mock_logger_fixture.info.assert_called_once_with(
        "File 'file.pdf' saved to '/fake/path/to/file.pdf'."
    )


@patch(OS_PATH_JOIN_PATH, return_value="/fake/path/to/file.pdf")
@patch(MOCK_OPEN_PATH, side_effect=IOError("Disk full"))
@patch(STREAMLIT_ERROR_PATH)
def test_save_uploaded_file_io_error(
    mock_st_error, mock_open_file, mock_os_join, mock_logger_fixture
):
    mock_uploaded_file = MagicMock()
    mock_uploaded_file.name = "file.pdf"
    mock_uploaded_file.getbuffer.return_value = b"file_content"

    result_path = save_uploaded_file(mock_uploaded_file)

    assert result_path is None
    mock_st_error.assert_called_once()
    mock_logger_fixture.error.assert_called_once()


# --- Tests for load_document (Adapted from test_rag_deep.py) ---


@patch(OS_PATH_BASENAME_PATH, return_value="sample.txt")
@patch(
    MOCK_OPEN_PATH,
    new_callable=mock_open,
    read_data="This is a test content from a TXT file.\nIt has multiple lines.",
)
def test_load_document_txt_success(mock_file, mock_basename, mock_logger_fixture):
    documents = load_document(SAMPLE_TXT_PATH)
    assert isinstance(documents, list)
    assert len(documents) == 1
    assert isinstance(documents[0], LangchainDocument)
    assert "This is a test content from a TXT file." in documents[0].page_content
    assert documents[0].metadata["source"] == "sample.txt"
    mock_logger_fixture.info.assert_called_with("Successfully loaded TXT: sample.txt")


@patch(OS_PATH_BASENAME_PATH, return_value="sample.docx")
@patch(DOCX_MODULE_PATH)  # Mock 'core.document_processing.docx'
def test_load_document_docx_success(
    mock_docx_module, mock_basename, mock_logger_fixture
):
    mock_doc_obj = MagicMock()
    mock_para1 = MagicMock()
    mock_para1.text = "This is test content from a DOCX file."
    mock_doc_obj.paragraphs = [mock_para1]
    mock_docx_module.Document.return_value = mock_doc_obj

    documents = load_document(SAMPLE_DOCX_PATH)
    assert isinstance(documents, list)
    assert len(documents) == 1
    assert isinstance(documents[0], LangchainDocument)
    assert "This is test content from a DOCX file." in documents[0].page_content
    assert documents[0].metadata["source"] == "sample.docx"
    mock_logger_fixture.info.assert_called_with("Successfully loaded DOCX: sample.docx")


@patch(OS_PATH_BASENAME_PATH, return_value="empty.txt")
@patch(MOCK_OPEN_PATH, new_callable=mock_open, read_data="")
@patch(STREAMLIT_WARNING_PATH)
def test_load_document_empty_txt(
    mock_st_warning, mock_file, mock_basename, mock_logger_fixture
):
    documents = load_document(EMPTY_TXT_PATH)
    assert documents == []
    mock_st_warning.assert_called_once_with(
        f"Text file '{os.path.basename(EMPTY_TXT_PATH)}' appears to be empty."
    )
    mock_logger_fixture.warning.assert_called_with(
        f"Text file '{os.path.basename(EMPTY_TXT_PATH)}' is empty."
    )


@patch(OS_PATH_BASENAME_PATH, return_value="empty.docx")
@patch(DOCX_MODULE_PATH)
@patch(STREAMLIT_WARNING_PATH)
def test_load_document_empty_docx(
    mock_st_warning, mock_docx_module, mock_basename, mock_logger_fixture
):
    mock_doc_obj = MagicMock()
    mock_doc_obj.paragraphs = []  # No paragraphs or paragraphs with no text
    mock_docx_module.Document.return_value = mock_doc_obj
    documents = load_document(EMPTY_DOCX_PATH)
    assert documents == []
    mock_st_warning.assert_called_once_with(
        f"DOCX file '{os.path.basename(EMPTY_DOCX_PATH)}' appears to be empty or contains no text."
    )
    mock_logger_fixture.warning.assert_called_with(
        f"DOCX file '{os.path.basename(EMPTY_DOCX_PATH)}' is empty or contains no text."
    )


@patch(OS_PATH_BASENAME_PATH, return_value="sample.png")
@patch(STREAMLIT_ERROR_PATH)
def test_load_document_unsupported_type(
    mock_st_error, mock_basename, mock_logger_fixture
):
    documents = load_document(UNSUPPORTED_FILE_PATH)  # .png is unsupported
    assert documents == []
    mock_st_error.assert_called_once_with(
        "Unsupported file type: '.png' for file 'sample.png'. Please upload a PDF, DOCX, or TXT file."
    )
    mock_logger_fixture.warning.assert_called_with(
        "Unsupported file type: '.png' for file 'sample.png'. Please upload a PDF, DOCX, or TXT file."
    )


@patch(OS_PATH_BASENAME_PATH, return_value="non_existent.txt")
@patch(MOCK_OPEN_PATH, side_effect=FileNotFoundError("File not found"))
@patch(
    STREAMLIT_ERROR_PATH
)  # load_document calls st.error in its generic exception handler
def test_load_document_non_existent_file(
    mock_st_error, mock_file_open, mock_basename, mock_logger_fixture
):
    documents = load_document(NON_EXISTENT_TXT_PATH)  # Path for a .txt file
    assert documents == []
    mock_st_error.assert_called_once()  # Check that st.error was called
    # More specific check for the logger message
    # FileNotFoundError is a subclass of IOError, which logs with logger.error in load_document
    mock_logger_fixture.error.assert_called_once()


@patch(OS_PATH_BASENAME_PATH, return_value="corrupted.docx")
@patch(DOCX_MODULE_PATH)  # Mock 'core.document_processing.docx'
@patch(STREAMLIT_ERROR_PATH)
def test_load_document_corrupted_docx(
    mock_st_error, mock_docx_module, mock_basename, mock_logger_fixture
):
    mock_docx_module.Document.side_effect = PackageNotFoundError(
        "Mocked error: Corrupted DOCX"
    )
    documents = load_document(CORRUPTED_DOCX_PATH)
    assert documents == []
    mock_st_error.assert_called_once_with(
        f"Failed to load DOCX '{os.path.basename(CORRUPTED_DOCX_PATH)}': The file appears to be corrupted or not a valid DOCX file."
    )
    mock_logger_fixture.error.assert_called_once()


@patch(OS_PATH_BASENAME_PATH, return_value="bad_encoding.txt")
@patch(MOCK_OPEN_PATH, new_callable=mock_open)
@patch(STREAMLIT_ERROR_PATH)
def test_load_document_txt_unicode_decode_error(
    mock_st_error, mock_file_open, mock_basename, mock_logger_fixture
):
    # Configure the mock_open instance to raise UnicodeDecodeError on read
    mock_file_instance = mock_file_open.return_value
    mock_file_instance.read.side_effect = UnicodeDecodeError(
        "utf-8", b"\x80", 0, 1, "invalid start byte"
    )

    documents = load_document(BAD_ENCODING_TXT_PATH)
    assert documents == []
    mock_st_error.assert_called_once_with(
        f"Failed to load TXT file '{os.path.basename(BAD_ENCODING_TXT_PATH)}': The file is not UTF-8 encoded. Please ensure it's a plain text file with UTF-8 encoding."
    )
    mock_logger_fixture.error.assert_called_once()


@patch(OS_PATH_BASENAME_PATH, return_value="corrupted.pdf")
@patch(PDF_PLUMBER_LOADER_PATH)
@patch(STREAMLIT_ERROR_PATH)
def test_load_document_pdf_syntax_error(
    mock_st_error, mock_pdf_loader, mock_basename, mock_logger_fixture
):
    mock_loader_instance = mock_pdf_loader.return_value
    mock_loader_instance.load.side_effect = PDFSyntaxError( # Use direct import
        "Mocked PDF Syntax Error"
    )
    documents = load_document(CORRUPTED_PDF_PATH)
    assert documents == []
    mock_st_error.assert_called_once_with(
        f"Failed to load PDF '{os.path.basename(CORRUPTED_PDF_PATH)}': The file may be corrupted or not a valid PDF."
    )
    mock_logger_fixture.error.assert_called_once()


# --- Tests for chunk_documents (from previous step, confirmed good) ---


@patch(RECURSIVE_SPLITTER_PATH)
def test_chunk_documents_success(mock_splitter_class, mock_logger_fixture):
    mock_splitter_instance = MagicMock()
    mock_chunk = LangchainDocument(
        page_content="chunked content", metadata={"source": "test.pdf"}
    )
    mock_splitter_instance.split_documents.return_value = [mock_chunk]
    mock_splitter_class.return_value = mock_splitter_instance

    raw_docs = [
        LangchainDocument(
            page_content="This is a long document.", metadata={"source": "test.pdf"}
        )
    ]
    chunks = chunk_documents(raw_docs)

    mock_splitter_class.assert_called_once_with(
        chunk_size=1000, chunk_overlap=200, add_start_index=True
    )
    mock_splitter_instance.split_documents.assert_called_once_with(raw_docs)
    assert len(chunks) == 1
    assert chunks[0].page_content == "chunked content"
    assert chunks[0].metadata["source"] == "test.pdf"
    mock_logger_fixture.info.assert_any_call("Starting to chunk 1 raw document(s).")
    mock_logger_fixture.info.assert_any_call("Chunking complete, 1 chunks created.")


@patch(STREAMLIT_WARNING_PATH)
def test_chunk_documents_empty_raw_docs(mock_st_warning, mock_logger_fixture):
    chunks = chunk_documents([])
    assert chunks == []
    mock_st_warning.assert_called_once_with(
        "No content found in the document to chunk."
    )
    mock_logger_fixture.warning.assert_called_once_with(
        "chunk_documents called with no raw documents."
    )


@patch(RECURSIVE_SPLITTER_PATH)
@patch(STREAMLIT_WARNING_PATH)
def test_chunk_documents_no_processable_chunks(
    mock_st_warning, mock_splitter_class, mock_logger_fixture
):
    mock_splitter_instance = MagicMock()
    mock_splitter_instance.split_documents.return_value = []
    mock_splitter_class.return_value = mock_splitter_instance

    raw_docs = [LangchainDocument(page_content=".")]
    chunks = chunk_documents(raw_docs)

    assert chunks == []
    mock_st_warning.assert_called_once_with(
        "Document chunking resulted in no processable text chunks. The document might be valid but contain no extractable text after initial processing, or the text is too short."
    )
    mock_logger_fixture.warning.assert_called_once_with(
        "Document chunking resulted in no processable text chunks."
    )


@patch(RECURSIVE_SPLITTER_PATH, side_effect=Exception("Chunking failed"))
@patch(STREAMLIT_ERROR_PATH)
def test_chunk_documents_exception(
    mock_st_error, mock_splitter_class_exception, mock_logger_fixture
):
    raw_docs = [LangchainDocument(page_content="Some content")]
    chunks = chunk_documents(raw_docs)

    assert chunks == []
    mock_st_error.assert_called_once()
    mock_logger_fixture.exception.assert_called_once()


# --- Tests for index_documents (from previous step, confirmed good) ---


@patch("core.document_processing.st.session_state", new_callable=MagicMock)
def test_index_documents_success(mock_session_state, mock_logger_fixture):
    mock_vector_db_instance = MagicMock()
    mock_session_state.DOCUMENT_VECTOR_DB = mock_vector_db_instance

    mock_chunks = [
        LangchainDocument(page_content="chunk1"),
        LangchainDocument(page_content="chunk2"),
    ]
    index_documents(mock_chunks)

    mock_vector_db_instance.add_documents.assert_called_once_with(mock_chunks)
    assert mock_session_state.document_processed is True
    mock_logger_fixture.info.assert_any_call(
        f"Indexing {len(mock_chunks)} document chunks."
    )
    mock_logger_fixture.info.assert_any_call(
        "Document chunks indexed successfully into vector store."
    )


@patch(STREAMLIT_WARNING_PATH)
def test_index_documents_empty_chunks(mock_st_warning, mock_logger_fixture):
    index_documents([])
    mock_st_warning.assert_called_once_with(
        "No document chunks available to index. This may happen if the document was empty or text extraction failed."
    )
    mock_logger_fixture.warning.assert_called_once_with(
        "index_documents called with no chunks to index."
    )


@patch("core.document_processing.st.session_state", new_callable=MagicMock)
@patch(STREAMLIT_ERROR_PATH)
def test_index_documents_exception_on_add(
    mock_st_error, mock_session_state, mock_logger_fixture
):
    mock_vector_db_instance = MagicMock()
    mock_vector_db_instance.add_documents.side_effect = Exception("DB Error")
    mock_session_state.DOCUMENT_VECTOR_DB = mock_vector_db_instance
    mock_session_state.document_processed = True

    mock_chunks = [LangchainDocument(page_content="chunk1")]
    index_documents(mock_chunks)

    mock_st_error.assert_called_once()
    mock_logger_fixture.exception.assert_called_once()
    assert mock_session_state.document_processed is False
